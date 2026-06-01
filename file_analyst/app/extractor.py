"""
Stage 1 — Extraction
Stage 2 — Structure Analysis

Soporta: .docx, .pdf, .txt, .xls/.xlsx, .csv, .json
Devuelve texto limpio + metadatos de estructura.
"""

import json
import csv
import re
import logging
from io import StringIO
from typing import Dict, Any

logger = logging.getLogger(__name__)


def extract_document(file_path: str, file_type: str) -> Dict[str, Any]:
    """
    Punto de entrada principal del extractor.
    Devuelve dict con: raw_text, structure, file_type
    """
    extractors = {
        "word":  _extract_docx,
        "pdf":   _extract_pdf,
        "text":  _extract_txt,
        "excel": _extract_excel,
        "csv":   _extract_csv,
        "json":  _extract_json,
    }

    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Extractor no disponible para tipo: {file_type}")

    logger.info(f"[Stage 1] Extrayendo {file_type}: {file_path}")
    raw_text, structure = extractor(file_path)

    # Limpieza universal
    clean_text = _clean_text(raw_text)

    logger.info(f"[Stage 2] Estructura: {structure.get('word_count', 0)} palabras, "
                f"{structure.get('sections', 0)} secciones")

    return {
        "raw_text": clean_text,
        "structure": structure,
        "file_type": file_type,
    }


# ──────────────────────────────────────────
# Extractores por tipo
# ──────────────────────────────────────────

def _extract_docx(path: str):
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("Instala python-docx: pip install python-docx")

    doc = Document(path)
    sections = []
    paragraphs = []
    tables_text = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        paragraphs.append(text)
        style = para.style.name if para.style else ""
        if "Heading" in style:
            sections.append({"heading": text, "level": style})

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        if rows:
            tables_text.append("\n".join(rows))

    full_text = "\n".join(paragraphs)
    if tables_text:
        full_text += "\n\n[TABLAS]\n" + "\n\n".join(tables_text)

    structure = {
        "sections": sections,
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables),
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
    }
    return full_text, structure


def _extract_pdf(path: str):
    try:
        import pdfplumber
    except ImportError:
        raise RuntimeError("Instala pdfplumber: pip install pdfplumber")

    pages_text = []
    total_tables = 0

    with pdfplumber.open(path) as pdf:
        num_pages = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages_text.append(text)
            tables = page.extract_tables()
            for table in (tables or []):
                rows = [" | ".join(str(c) for c in row if c) for row in table if any(row)]
                if rows:
                    pages_text.append("[TABLA]\n" + "\n".join(rows))
                    total_tables += 1

    full_text = "\n".join(pages_text)
    structure = {
        "page_count": num_pages,
        "table_count": total_tables,
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
        "sections": [],
    }
    return full_text, structure


def _extract_txt(path: str):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()

    lines = [l.strip() for l in text.splitlines() if l.strip()]
    structure = {
        "line_count": len(lines),
        "word_count": len(text.split()),
        "char_count": len(text),
        "sections": [],
    }
    return text, structure


def _extract_excel(path: str):
    try:
        import openpyxl
    except ImportError:
        raise RuntimeError("Instala openpyxl: pip install openpyxl")

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets_text = []
    sheet_info = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                rows_text.append(" | ".join(cells))
        if rows_text:
            sheets_text.append(f"[HOJA: {sheet_name}]\n" + "\n".join(rows_text))
            sheet_info.append({"sheet": sheet_name, "rows": len(rows_text)})

    full_text = "\n\n".join(sheets_text)
    structure = {
        "sheet_count": len(wb.sheetnames),
        "sheets": sheet_info,
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
        "sections": [],
    }
    return full_text, structure


def _extract_csv(path: str):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample)
        except csv.Error:
            dialect = csv.excel

        reader = csv.reader(f, dialect)
        rows = []
        headers = []
        for i, row in enumerate(reader):
            if i == 0:
                headers = row
            rows.append(" | ".join(row))

    full_text = "\n".join(rows)
    structure = {
        "row_count": len(rows),
        "column_count": len(headers),
        "headers": headers,
        "word_count": len(full_text.split()),
        "char_count": len(full_text),
        "sections": [],
    }
    return full_text, structure


def _extract_json(path: str):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        data = json.load(f)

    pretty = json.dumps(data, ensure_ascii=False, indent=2)

    # Resumen plano para análisis
    if isinstance(data, list):
        summary_text = f"Array JSON con {len(data)} elementos.\n" + pretty[:8000]
    elif isinstance(data, dict):
        keys = list(data.keys())
        summary_text = f"Objeto JSON con claves: {', '.join(keys[:20])}.\n" + pretty[:8000]
    else:
        summary_text = pretty[:8000]

    structure = {
        "type": type(data).__name__,
        "keys": list(data.keys()) if isinstance(data, dict) else [],
        "length": len(data) if isinstance(data, (list, dict)) else 1,
        "word_count": len(summary_text.split()),
        "char_count": len(summary_text),
        "sections": [],
    }
    return summary_text, structure


# ──────────────────────────────────────────
# Limpieza universal de texto
# ──────────────────────────────────────────

def _clean_text(text: str) -> str:
    """
    Normaliza el texto extraído:
    - Elimina líneas vacías excesivas
    - Elimina caracteres de control
    - Normaliza espacios
    """
    # Eliminar caracteres de control (excepto \n y \t)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    # Normalizar múltiples espacios
    text = re.sub(r"[ \t]+", " ", text)
    # Normalizar múltiples saltos de línea
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Eliminar líneas que solo tienen espacios
    lines = [l if l.strip() else "" for l in text.splitlines()]
    text = "\n".join(lines).strip()
    return text
